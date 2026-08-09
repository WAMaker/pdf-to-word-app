# -*- coding: utf-8 -*-
"""
PDF 转 Word 核心转换模块
-------------------------
技术要点(解决"改字号后异常换行/缩进"问题):
1. 段落重建:PDF 中同一逻辑段落的分散行会被合并为真正的 Word 段落,
   而不是每行一个孤立段落,避免改字号后 Word 重新排版时乱换行。
2. 格式统一:所有正文统一使用 Word 的"正文"样式,字号由样式控制,
   不保留 PDF 的逐行硬编码格式(行距/缩进),改字号时全文自动重排。
3. 合理缩进:首行缩进用段落属性实现(而不是手动空格),改字号不会错位。

依赖:PyMuPDF(fitz)、python-docx
"""
from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------

@dataclass
class TextLine:
    """PDF 中提取的一行文本(可能只是段落的一部分)"""
    text: str
    bbox: Tuple[float, float, float, float]  # x0, y0, x1, y1
    size: float          # 字号(pt)
    font: str
    bold: bool
    align: Optional[str]  # 'left' | 'center' | 'right' | 'justify' | None
    spans: List[Tuple[str, bool]] = field(default_factory=list)  # (文本, 加粗) 行内分段样式

    def __post_init__(self):
        # 若无 span 级信息,回退为整行一条
        if not self.spans and self.text:
            self.spans = [(self.text, self.bold)]


@dataclass
class ImageBlock:
    """PDF 中提取的图片块"""
    bbox: Tuple[float, float, float, float]
    xref: int             # PDF 对象编号,用于提取图片数据
    width: int            # 像素宽
    height: int           # 像素高
    page_index: int = 0


@dataclass
class Paragraph:
    """重建后的逻辑段落(可能是文本段落或图片)"""
    lines: List[TextLine] = field(default_factory=list)
    image: Optional[ImageBlock] = None  # 图片段落时非空
    style: str = "body"      # 'title' | 'heading' | 'body' | 'caption' | 'image'
    align: str = "left"
    indent_first: float = 0.0   # 首行缩进(pt)
    indent_left: float = 0.0    # 左缩进(pt)
    bold: bool = False          # 段落是否加粗(PDF 原样保留)

    @property
    def text(self) -> str:
        if self.image is not None:
            return "[图片]"
        # 合并行内文本:英文单词间需要空格,中文直接拼接
        parts = []
        for line in self.lines:
            parts.append(line.text)
        return "".join(parts) if _is_cjk_text("".join(parts)) else " ".join(p.strip() for p in parts)

    @property
    def size(self) -> float:
        if not self.lines:
            return 12.0
        return max(l.size for l in self.lines)


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------

_CJK_RE = re.compile(r"[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]")


def _is_bold_font(font_name: str) -> bool:
    """判断字体是否为粗体(PDF 结构化属性:字体名含 Bold/Black/Heavy)"""
    if not font_name:
        return False
    name = font_name.lower()
    return any(k in name for k in ("bold", "black", "heavy", "demibold", "semibold"))


def _is_cjk_text(text: str) -> bool:
    """判断文本是否以中文为主"""
    if not text:
        return False
    cjk = len(_CJK_RE.findall(text))
    return cjk / max(len(text), 1) > 0.3


_LIST_MARKER_RE = re.compile(r"^(?:[l1iI•·▪●]|[0-9]{1,2}[.、.)])\s+")


def _clean_text(s: str) -> str:
    """清理文本:去空白、修正常见 OCR/排版残留、剔除非法 XML 控制字符"""
    # 剔除 XML 1.0 不允许的控制字符(NULL、0x01-0x08、0x0B、0x0C、0x0E-0x1F 等)
    s = re.sub(
        r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]",
        "",
        s,
    )
    s = s.replace("\u3000", " ").replace("\xa0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = s.strip()
    # 清理行首列表符残留(PDF 中项目符号常被渲染为 'l'/'1'/'•')
    s = _LIST_MARKER_RE.sub("", s)
    # 单独一个列表符字符(无后续文字)直接清空
    if re.fullmatch(r"[l1iI•·▪●]", s):
        return ""
    return s


def _is_blank_line(s: str) -> bool:
    return not s.strip()


# ---------------------------------------------------------------------------
# 段落重建
# ---------------------------------------------------------------------------

def _extract_images(page: fitz.Page) -> List[ImageBlock]:
    """提取页面中的图片块"""
    images: List[ImageBlock] = []
    page_images = page.get_images(full=True)
    seen = set()
    for img in page_images:
        xref = img[0]
        if xref in seen:
            continue
        seen.add(xref)
        try:
            rects = page.get_image_rects(xref)
        except Exception:  # noqa: BLE001
            continue
        if not rects:
            continue
        # 取最大的一块显示区域
        rect = max(rects, key=lambda r: (r.width * r.height))
        # 过滤:占满整页的背景图(如扫描件)跳过,避免把扫描页当插图
        page_area = page.rect.width * page.rect.height
        if rect.width * rect.height > page_area * 0.9:
            continue
        images.append(ImageBlock(
            bbox=(rect.x0, rect.y0, rect.x1, rect.y1),
            xref=xref,
            width=img[2],
            height=img[3],
        ))
    return images


def _extract_elements(page: fitz.Page) -> Tuple[List[TextLine], List[ImageBlock]]:
    """同时提取文本行和图片(按 y 坐标排序的元素流由调用方处理)"""
    lines = _extract_lines(page)
    images = _extract_images(page)
    return lines, images


def _merge_elements(lines: List[TextLine], images: List[ImageBlock]):
    """把文本行和图片块按垂直位置合并成有序元素流
    返回 (y, kind, obj) 元组列表:kind='line'|'image'
    """
    elements = []
    for line in lines:
        elements.append((line.bbox[1], "line", line))
    for img in images:
        elements.append((img.bbox[1], "image", img))
    elements.sort(key=lambda e: (e[0], 0 if e[1] == "line" else 1))
    return elements


def _extract_lines(page: fitz.Page) -> List[TextLine]:
    """从页面提取文本行(使用 dict 模式获得字体/字号/位置信息)"""
    lines: List[TextLine] = []
    blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
    for block in blocks:
        if block.get("type") != 0:  # 只处理文本块,跳过图片
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            raw = "".join(s.get("text", "") for s in spans)
            text = _clean_text(raw)
            if not text:
                continue
            bbox = line.get("bbox", (0, 0, 0, 0))
            sizes = [s.get("size", 12.0) for s in spans]
            fonts = [s.get("font", "") for s in spans]
            # 加粗 = 字体名含 Bold/Black/Heavy(PDF 结构化样式属性)
            # 这是 PDF 自身的字体信息,比 flags 位更可靠,也避免内容特判
            bold = any(_is_bold_font(f) for f in fonts)
            # 行内分段样式:逐 span 收集 (文本, 加粗),供行内部分加粗保留
            span_styles: List[Tuple[str, bool]] = []
            for s in spans:
                st = _clean_text(s.get("text", ""))
                if not st:
                    continue
                span_styles.append((st, _is_bold_font(s.get("font", ""))))
            if not span_styles:
                span_styles = [(text, bold)]
            lines.append(TextLine(
                text=text,
                bbox=bbox,
                size=max(sizes),
                font=max(fonts, key=fonts.count),
                bold=bold,
                align=_detect_align(line, page.rect.width),
                spans=span_styles,
            ))
    return lines


def _detect_align(line, page_width: float) -> str:
    x0, _, x1, _ = line.get("bbox", (0, 0, 0, 0))
    line_width = x1 - x0
    left_margin = x0
    right_margin = page_width - x1
    width_ratio = line_width / page_width
    # 满宽行(占页宽 55% 以上):两端对齐或左对齐正文,不是居中
    # (首行缩进段落的段首行也常达到 ~60% 宽,不能判居中)
    if width_ratio > 0.55:
        if left_margin < 8 and right_margin < 8:
            return "justify"
        return "left"
    # 短行(≤55%):居中需同时满足:左右边距都较明显(>30pt)且接近
    if left_margin > 30 and right_margin > 30 \
            and abs(left_margin - right_margin) <= max(12, min(left_margin, right_margin) * 0.35):
        return "center"
    if left_margin > right_margin * 2.2 and left_margin > 20:
        return "right"
    return "left"


def _infer_style(line: TextLine, is_first: bool, para_lines: List[TextLine],
                 body_size: float = 0.0) -> str:
    """推断段落样式:标题 / 小标题 / 正文(启发式路径用)
    策略:只使用 PDF 结构化样式信号——相对字号(与正文字号比较)、
    加粗(字体名)、单行长度、段落结构;不做内容特判。
    :param body_size: 页面正文字号估计(0 表示未知,退回绝对阈值)
    :param is_first: 是否全文档第一个非空文本段落(课程/文档大标题通常在此)
    """
    para_text = "".join(l.text for l in para_lines)
    para_len = len(para_text.strip())
    # 以句末标点结尾的通常是正文/完整句子,不是标题
    ends_with_punct = para_text.rstrip().endswith((
        "。", ".", "！", "？", "!", "?", "，", ",", "；", ";", "、", "：", ":",
    ))

    # 全文档首段 + 加粗 + 单行 + 短文本 → 大标题(如课程标题)
    if is_first and line.bold and len(para_lines) == 1 and para_len <= 40:
        return "title"

    # 绝对阈值兜底:超大字号必为标题
    if line.size >= 20:
        return "title"
    if body_size > 0:
        ratio = line.size / body_size
        # 明显大(≥1.3x)且段落短(<60字)→ 标题
        if ratio >= 1.3 and para_len < 60 and not ends_with_punct:
            return "title" if ratio >= 1.8 else "heading"
        # 略大(≥1.05x)且加粗且段落很短(≤20字)→ 小标题
        if ratio >= 1.05 and line.bold and para_len <= 20 and not ends_with_punct:
            return "heading"
        # 同字号加粗且段落短(≤18字)且单行且不以标点结尾 → 小标题
        if abs(ratio - 1.0) < 0.05 and line.bold and para_len <= 18 \
                and len(para_lines) == 1 and not ends_with_punct:
            return "heading"
        return "body"
    # 未知正文字号时:字号 >= 14 且加粗/居中/短文本 → 标题
    if line.size >= 14:
        if (line.bold or line.align == "center" or para_len < 30) and not ends_with_punct:
            return "heading"
    return "body"


def _estimate_body_size(lines: List[TextLine]) -> float:
    """估计页面正文字号:取非加粗文本量最大的字号(众数加权)
    排除加粗行,避免标题/强调行干扰正文基准。
    """
    if not lines:
        return 0.0
    from collections import Counter
    counter = Counter()
    for line in lines:
        if line.text and not line.bold:
            counter[round(line.size, 1)] += len(line.text)
    if not counter:
        # 全部加粗时退回全量统计
        for line in lines:
            if line.text:
                counter[round(line.size, 1)] += len(line.text)
    if not counter:
        return 0.0
    return counter.most_common(1)[0][0]


def _same_paragraph(prev: TextLine, cur: TextLine, para_gap: float, line_gap: float) -> bool:
    """判断两行是否属于同一逻辑段落(启发式路径,仅用于非 Tagged PDF)
    :param para_gap: 段落间距阈值(动态计算,pt)
    :param line_gap: 行内间距基准(pt,取中位数)
    段落边界识别只使用 PDF 结构化信号(布局/样式属性),不做内容特判:
    1. 行距显著大于行内基准 → 新段(真实的段间距)
    2. 当前行比上一行明显右移(>12pt)→ 新段(首行缩进,最常见的段界标志)
    3. 加粗→不加粗 且 上一行明显短(<60%当前行宽)→ 新段(标题行特征)
    4. 其余情况合并为同一段。
    注意:不依赖 PDF 文本块 ID(同一段经常被拆成多个 block),
    也不在浮点边缘比较行距(有的行距 16.7、有的 16.8,阈值取中位数会误判)。
    """
    # 垂直距离:当前行顶部 - 上一行底部
    gap = cur.bbox[1] - prev.bbox[3]

    # 1) 行距明显大于行内基准(>1.6x)→ 新段落
    if gap > para_gap:
        return False

    # 2) 当前行有显著左缩进(>12pt),通常是新段落(如正文首行缩进)
    if cur.bbox[0] - prev.bbox[0] > 12:
        return False

    # 3) 加粗短行(标题行特征):上一行加粗且明显比当前行短(<60%)→ 新段
    # 注意:长行加粗→不粗多为段内强调(如穴位名),不在此列
    if prev.bold:
        prev_width = prev.bbox[2] - prev.bbox[0]
        cur_width = cur.bbox[2] - cur.bbox[0]
        if prev_width > 0 and prev_width < cur_width * 0.6:
            return False

    # 4) 默认:同一段落
    return True


def _estimate_para_gap(lines: List[TextLine]) -> Tuple[float, float]:
    """动态估算间距阈值
    :return: (段落阈值, 行内间距基准)
    行内间距取低分位值(20%),段落阈值 = 行内间距 * 2.0(至少 10pt)。
    先按 y 坐标排序再统计,避免 PDF 文本块顺序(非视觉顺序)污染行距统计;
    低分位避免被标题/图片间的大间距拉高阈值。
    注意:很多 PDF 段间距 == 行间距(如本工具样本文档),此时段界靠首行缩进判断。
    """
    sorted_lines = sorted(lines, key=lambda l: l.bbox[1])
    gaps = []
    for i in range(1, len(sorted_lines)):
        g = sorted_lines[i].bbox[1] - sorted_lines[i - 1].bbox[3]
        if g > 0:
            gaps.append(g)
    if not gaps:
        return 20.0, 4.0
    gaps.sort()
    # 20% 分位作为行内典型间距
    idx = min(len(gaps) - 1, max(0, int(len(gaps) * 0.2)))
    line_gap = gaps[idx]
    return max(line_gap * 2.0, 10.0), line_gap


def _align_from_lines(para: Paragraph, page_width: float) -> str:
    """根据行位置综合判断对齐方式"""
    if not para.lines:
        return "left"
    aligns = [l.align for l in para.lines]
    from collections import Counter
    return Counter(aligns).most_common(1)[0][0] or "left"


def _infer_para_bold(lines: List[TextLine]) -> bool:
    """推断段落是否加粗:多数行加粗(≥50%)则段落加粗
    保留 PDF 原文的加粗样式(强调、章节说明等)。
    """
    if not lines:
        return False
    bold_count = sum(1 for l in lines if l.bold)
    return bold_count / len(lines) >= 0.5


def rebuild_paragraphs(lines: List[TextLine], page_width: float, page_height: float,
                       images: Optional[List[ImageBlock]] = None,
                       first_text: bool = True) -> List[Paragraph]:
    """把提取的行重建为逻辑段落(支持图片穿插)
    :param first_text: 本次调用中第一个文本段落是否属于全文档第一个文本段落
                       (用于标题识别;多页文档时跨页传递)
    """
    paragraphs: List[Paragraph] = []
    current: Optional[Paragraph] = None

    elements = _merge_elements(lines, images or [])
    para_gap, line_gap = _estimate_para_gap(lines)

    # 页面内容区左边缘(即 PDF 页面左边距),用于计算相对缩进
    page_left = min((l.bbox[0] for l in lines), default=0.0)

    for _, kind, obj in elements:
        if kind == "image":
            # 图片:先结束当前文本段落,再作为独立段落插入
            if current is not None:
                paragraphs.append(current)
                current = None
            paragraphs.append(Paragraph(
                image=obj,
                style="image",
                align="center",
            ))
            continue

        line = obj  # TextLine
        if _is_blank_line(line.text):
            if current is not None:
                paragraphs.append(current)
                current = None
            continue

        if current is None:
            current = Paragraph(lines=[line])
            current.align = line.align
            continue

        if _same_paragraph(current.lines[-1], line, para_gap, line_gap):
            current.lines.append(line)
        else:
            paragraphs.append(current)
            current = Paragraph(lines=[line])
            current.align = line.align

    if current is not None:
        paragraphs.append(current)

    # 后处理:样式推断 + 对齐修正(跳过图片段落)
    body_size = _estimate_body_size(lines)
    is_first_local = first_text
    for para in paragraphs:
        if para.image is not None:
            continue
        para.align = _align_from_lines(para, page_width)
        para.style = _infer_style(para.lines[0], is_first_local, para.lines, body_size)
        para.bold = _infer_para_bold(para.lines)
        is_first_local = False  # 仅全文档第一个文本段落享受 title 候选
        # 首行缩进检测:第二行比第一行靠左,且第一行有缩进 → 段落缩进
        if len(para.lines) > 1:
            first_x0 = para.lines[0].bbox[0]
            rest_x0 = min(l.bbox[0] for l in para.lines[1:])
            indent = first_x0 - rest_x0
            if indent > 8:
                para.indent_first = indent
            # 左缩进 = 相对页面左边距的额外缩进(PDF 绝对 x0 含页面边距,
            # 直接写入 Word 会导致正文整体严重偏右)
            extra_left = rest_x0 - page_left
            if extra_left > 8:
                para.indent_left = extra_left

    return paragraphs


# ---------------------------------------------------------------------------
# 结构树解析(Tagged PDF 官方段落边界)
# ---------------------------------------------------------------------------

_ACTUALTEXT_CLEAN = {"(": " ", ")": " ", "（": " ", "）": " ", "( ": " ", " )": " "}


def _decode_actualtext(raw: str) -> str:
    """解码 ActualText:<FEFF...> 是 UTF-16BE hex,否则是字面量"""
    raw = raw.strip()
    if raw.startswith("<FEFF") and raw.endswith(">"):
        try:
            return bytes.fromhex(raw[5:-1]).decode("utf-16-be", errors="replace")
        except Exception:  # noqa: BLE001
            return ""
    if raw.startswith("<") and raw.endswith(">"):
        try:
            hexs = raw[1:-1]
            if len(hexs) % 4 == 0:
                return bytes.fromhex(hexs).decode("utf-16-be", errors="replace")
            return bytes.fromhex(hexs).decode("latin1", errors="replace")
        except Exception:  # noqa: BLE001
            return ""
    # 字面量:过滤孤立的括号残留(Word 导出 PDF 时空格的错误表示)
    return _ACTUALTEXT_CLEAN.get(raw, raw)


@dataclass
class StructPara:
    """结构树段落(Tagged PDF 作者真实段落)"""
    text: str
    type: str           # 'P' | 'H' | 'H1'..'H3' | 'LI' | 'Title' | ...
    page_index: int     # 0-based
    bold: bool = False  # 后续由行信息补充
    y_pos: float = 0.0  # 段落 y 位置(用于图片插入定位)
    x_pos: float = 0.0  # 段落 x 位置(首行缩进推断)


def _norm_for_match(s: str) -> str:
    """规范化文本用于行匹配:去括号/方括号/空格
    结构树 ActualText 常带 Word 导出括号(如 '(YS05)'、'(3)'、'(])'),
    页面行文本没有;统一去括号去空格后再匹配。
    """
    s = re.sub(r"[\(\)（）\[\]]", "", s)
    s = re.sub(r"\s+", "", s)
    return s


def _extract_structured_paragraphs(pdf: fitz.Document) -> Optional[List[StructPara]]:
    """从 Tagged PDF 结构树提取作者真实段落。
    返回 None 表示该 PDF 无结构树(退回启发式段落重建)。
    递归遍历结构树,支持:
    - P/H1-H3/Title → 段落
    - L(列表容器) → 递归进 LI(列表项)
    - LI/LBody → 列表项段落
    - Span → 文本片段(ActualText)
    完全依赖结构标签,不做内容特判。
    """
    # 1. 找 StructTreeRoot
    root = None
    for xref in range(1, pdf.xref_length()):
        try:
            obj = pdf.xref_object(xref, compressed=False)
        except Exception:  # noqa: BLE001
            continue
        if "/Type /Catalog" in obj and "/StructTreeRoot" in obj:
            m = re.search(r"/StructTreeRoot\s+(\d+)\s+0\s+R", obj)
            if m:
                root = int(m.group(1))
                break
    if not root:
        return None
    try:
        root_obj = pdf.xref_object(root, compressed=False)
        m = re.search(r"/K\s*\[\s*(\d+)\s+0\s+R", root_obj)
        if not m:
            return None
        doc_elem = int(m.group(1))
    except Exception:  # noqa: BLE001
        return None

    paras: List[StructPara] = []

    def _page_of(page_xref: Optional[int]) -> int:
        if not page_xref:
            return 0
        for pi in range(pdf.page_count):
            if pdf[pi].xref == page_xref:
                return pi
        return 0

    # 每页行缓存(用于 y/x 定位和加粗匹配)
    page_lines_cache: dict = {}
    for pi in range(pdf.page_count):
        page_lines_cache[pi] = _extract_lines(pdf[pi])

    def _elem_text(elem_xref: int, depth: int = 0) -> str:
        """递归提取元素文本(聚合所有后代 Span 的 ActualText)"""
        if depth > 15:
            return ""
        try:
            obj = pdf.xref_object(elem_xref, compressed=False)
        except Exception:  # noqa: BLE001
            return ""
        if "/StructElem" not in obj:
            return ""
        sm = re.search(r"/S\s+/(\w+)", obj)
        stype = sm.group(1) if sm else "?"
        text = ""
        if stype == "Span":
            atext_m = re.search(r"/ActualText\s+([^\s/]+)", obj)
            if atext_m:
                text += _decode_actualtext(atext_m.group(1))
        # 递归子节点:只取 /K 数组中的引用(排除 /P 父引用、/Pg 页面引用)
        kids = []
        km = re.search(r"/K\s*\[([^\]]*)\]", obj)
        if km:
            kids = [int(x) for x in re.findall(r"(\d+)\s+0\s+R", km.group(1))]
        for sk in kids:
            if sk == elem_xref:
                continue
            text += _elem_text(sk, depth + 1)
        return text

    def _walk(elem_xref: int, depth: int = 0):
        """递归遍历结构树,输出段落"""
        if depth > 15:
            return
        try:
            obj = pdf.xref_object(elem_xref, compressed=False)
        except Exception:  # noqa: BLE001
            return
        if "/StructElem" not in obj:
            return
        sm = re.search(r"/S\s+/(\w+)", obj)
        stype = sm.group(1) if sm else "?"
        pgm = re.search(r"/Pg\s+(\d+)\s+0\s+R", obj)
        page_xref = int(pgm.group(1)) if pgm else None
        page_index = _page_of(page_xref)

        # 段落类型:输出并聚合文本
        is_para_type = stype in ("P", "H", "H1", "H2", "H3", "Title", "LI", "LBody")
        if is_para_type:
            text = _elem_text(elem_xref, depth + 1)
            text = re.sub(r"\s*\n\s*", "", text)
            text = re.sub(r"[ \t]{2,}", " ", text).strip()
            # 清理 Word 导出残留括号(圆括号为导出标记;方括号是原文题目标记,保留)
            text = re.sub(r"[\(\)（）]", "", text)
            text = _LIST_MARKER_RE.sub("", text)  # 清理列表符残留(如 'l 胃經腹部...')
            text = text.strip()
            if text:
                sp2 = StructPara(text=text, type=stype, page_index=page_index)
                # 定位 y/x:匹配该页行文本(规范化)
                sp_norm = _norm_for_match(text)
                for line in page_lines_cache.get(page_index, []):
                    ln = _norm_for_match(line.text)
                    if not ln:
                        continue
                    if ln in sp_norm or sp_norm in ln:
                        sp2.y_pos = line.bbox[1]
                        sp2.x_pos = max(sp2.x_pos, line.bbox[0])
                        break
                paras.append(sp2)
            # 段落元素文本已聚合,不递归输出子元素(避免 LBody 重复)
            return

        # 容器类型(L 列表/Sect 等):递归子节点(只取 /K 数组引用)
        kids = []
        km = re.search(r"/K\s*\[([^\]]*)\]", obj)
        if km:
            kids = [int(x) for x in re.findall(r"(\d+)\s+0\s+R", km.group(1))]
        for sk in kids:
            if sk == elem_xref:
                continue
            _walk(sk, depth + 1)

    _walk(doc_elem)

    if not paras:
        return None
    # 按页排序(保持文档顺序)
    paras.sort(key=lambda p: p.page_index)
    return paras


# ---------------------------------------------------------------------------
# DOCX 生成
# ---------------------------------------------------------------------------

FONT_SIZE_MAP = {
    "小四": 12,
    "四号": 14,
    "小三": 15,
    "三号": 16,
    "小二": 18,
    "二号": 22,
}


def _setup_page(doc: Document):
    """设置 A4 页面(中文文档标准),边距适中
    Word 默认是 Letter(215.9x279.4mm),与中文 PDF 的 A4 容量不同;
    不匹配会导致固定分页符下"一页没占满就换页"或内容溢出。
    """
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(25.4)   # 1 英寸
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)


def _setup_style(doc: Document, font_name: str, font_size_pt: int):
    """设置文档默认样式,统一控制字体字号"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _extract_image_bytes(pdf: fitz.Document, xref: int) -> Tuple[Optional[bytes], str]:
    """从 PDF 提取图片原始字节和扩展名"""
    try:
        info = pdf.extract_image(xref)
        if not info:
            return None, "png"
        return info["image"], info["ext"]
    except Exception:  # noqa: BLE001
        return None, "png"


def _add_image_to_doc(doc: Document, pdf: fitz.Document, img: ImageBlock,
                      image_dir: str, page_width: float) -> Optional[str]:
    """提取 PDF 图片并插入 docx,返回图片本地路径(供预览)
    :param image_dir: 图片缓存目录(预览用)
    """
    data, ext = _extract_image_bytes(pdf, img.xref)
    if not data:
        return None

    # 保存到缓存目录供预览使用
    os.makedirs(image_dir, exist_ok=True)
    img_path = os.path.join(image_dir, f"img_{img.xref}.{ext}")
    with open(img_path, "wb") as f:
        f.write(data)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = p.add_run()
        run.add_picture(img_path, width=Pt(min(img.bbox[2] - img.bbox[0], page_width * 0.9)))
    except Exception:  # noqa: BLE001
        # 图片格式 docx 不支持时跳过
        return None
    return img_path


def add_paragraph(doc: Document, para: Paragraph, font_name: str, base_size: float,
                 pdf: Optional[fitz.Document] = None,
                 image_dir: Optional[str] = None,
                 page_width: float = 595.0,
                 struct_path: bool = False) -> Optional[str]:
    """把重建段落写入 docx;图片段落返回本地图片路径
    :param struct_path: 结构树模式(标题不放大字号,忠实 PDF 样式)
    :return: 图片段落返回图片路径,文本段落返回 None
    """
    if para.image is not None and pdf is not None:
        return _add_image_to_doc(doc, pdf, para.image, image_dir or "", page_width)

    p = doc.add_paragraph()
    pf = p.paragraph_format

    # 对齐
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = align_map.get(para.align, WD_ALIGN_PARAGRAPH.LEFT)

    # 缩进用段落属性(不是空格)→ 改字号不错位
    if para.indent_first > 0:
        pf.first_line_indent = Pt(para.indent_first)
    if para.indent_left > 0:
        pf.left_indent = Pt(para.indent_left)

    # 样式相关字号
    # 结构树模式(struct_path=True)忠实 PDF 样式:标题不放大字号,
    # 加粗由字体决定;启发式模式保留标题放大(老行为)
    if para.style == "title":
        size = base_size if struct_path else max(base_size * 1.8, 26)
        bold = True
    elif para.style == "heading":
        size = base_size if struct_path else max(base_size * 1.4, 18)
        bold = True
    else:
        size = base_size

    # 写入文本:span 级加粗(行内部分加粗),保留 PDF 原文逐 span 的加粗/常规样式
    # 而不是整行/整段统一加粗——避免'原文只加粗穴位名,转换后整行变粗'的错判
    if para.style == "body":
        for line in para.lines:
            for span_text, span_bold in line.spans:
                run = p.add_run(span_text)
                _set_run_font(run, font_name, size, span_bold)
    else:
        text = para.text
        run = p.add_run(text)
        _set_run_font(run, font_name, size, bold)

    return None


_SENT_END_RE = re.compile(r"[。！？!?；;：:\”’」』）)]$")

def _is_para_complete(para: Paragraph) -> bool:
    """判断段落是否以句末标点结束(完整段落)。
    以句号/感叹号/问号/冒号/引号等结尾 → 完整;
    标题/小标题样式(如 '胃經腹部重點穴位診斷作用及功能複習')虽无标点,
    但作为独立标题也是完整的,不参与跨页续接。
    """
    if para.image is not None:
        return True
    if para.style in ("title", "heading"):
        return True
    if not para.lines:
        return True
    text = para.text.rstrip()
    if not text:
        return True
    return bool(_SENT_END_RE.search(text))


def convert_pdf_to_docx(
    pdf_path: str,
    docx_path: str,
    font_size_label: str = "三号",
    font_name: str = "微软雅黑",
    progress_cb=None,
    image_dir: Optional[str] = None,
    page_breaks: bool = False,
) -> dict:
    """
    转换主入口
    :param pdf_path: 输入 PDF
    :param docx_path: 输出 DOCX
    :param font_size_label: 字号('小四'/'四号'/'小三'/'三号'/'小二'/'二号')
    :param font_name: 中文字体名
    :param progress_cb: 进度回调(page_index, total_pages)
    :param image_dir: 图片缓存目录(用于预览;为 None 时不保留预览图片)
    :param page_breaks: 是否按 PDF 原页插入分页符。
        默认 False:Word 流式自动分页,改字号后全文重排,不会出现
        "一页没占满就换页"或内容溢出的半空页。
        设为 True 时保留 PDF 每页的分页结构(但改字号后可能留白/溢出)。
    :return: 转换结果,含 preview 结构化数据(段落+图片)
    """
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")

    base_size = FONT_SIZE_MAP.get(font_size_label, 16)
    doc = Document()
    _setup_page(doc)
    _setup_style(doc, font_name, base_size)

    pdf = fitz.open(pdf_path)
    total = pdf.page_count
    para_count = 0
    image_count = 0
    preview_blocks: List[dict] = []  # 预览结构化数据
    first_text_done = False  # 跨页追踪全文档第一个文本段落(标题识别用)

    # 首选:Tagged PDF 结构树(作者真实段落,无需启发式)
    struct_paras = _extract_structured_paragraphs(pdf)
    used_struct = struct_paras is not None

    try:
        if struct_paras:
            # 按页分组写入结构段落
            import itertools as _it
            by_page = {}
            for sp in struct_paras:
                by_page.setdefault(sp.page_index, []).append(sp)
            # 预提取每页行级加粗信息(结构树无 bold,需从行信息匹配)
            page_lines_cache = {}
            for i in range(pdf.page_count):
                page_lines_cache[i] = _extract_lines(pdf[i])
            # 预提取每页图片(y 定位用)
            page_imgs_cache = {}
            for i in range(pdf.page_count):
                page_imgs_cache[i] = _extract_images(pdf[i])
            struct_first_done = False  # 结构树路径跨页追踪全文档首个段落
            for i, page in enumerate(pdf):
                if progress_cb:
                    progress_cb(i + 1, total)
                if i > 0 and page_breaks:
                    doc.add_page_break()
                page_lines = page_lines_cache.get(i, [])
                # 构造本页输出项:段落 + 图片,按 y 合并排序
                items: List[tuple] = []  # (y, 'para'|'img', obj)
                for sp in by_page.get(i, []):
                    items.append((sp.y_pos, "para", sp))
                for img in page_imgs_cache.get(i, []):
                    items.append((img.bbox[1], "img", img))
                items.sort(key=lambda x: (x[0], 0 if x[1] == "para" else 1))
                for _, kind, obj in items:
                    if kind == "img":
                        # 图片:提取并插入当前位置
                        img_path = _add_image_to_doc(
                            doc, pdf, obj, image_dir or "", page.rect.width)
                        if img_path:
                            image_count += 1
                            preview_blocks.append({"type": "image", "path": img_path})
                        continue
                    sp = obj  # StructPara
                    sp_text = sp.text.strip()
                    # 样式推断:只用结构树标签(PDF 结构化信息)
                    # P/LI → 正文;H/H1/Title → 大标题;H2/H3 → 小标题
                    # 不做内容特判(如选项/答案/列表项),也不猜标题
                    if sp.type in ("H", "H1", "Title"):
                        style = "title"
                    elif sp.type in ("H2", "H3"):
                        style = "heading"
                    else:
                        style = "body"
                    para = Paragraph(lines=[], style=style, align="left")
                    # span 级加粗:按页面行的 span 片段在原文中顺序匹配
                    # 页面行文本与结构文本一致(无括号差异),直接 find
                    span_data: List[Tuple[str, bool]] = []
                    search_pos = 0
                    matched_any = False
                    for line in page_lines:
                        for span_text, span_bold in line.spans:
                            st = span_text.strip()
                            if not st or len(st) < 2:
                                continue
                            idx = sp_text.find(st, search_pos)
                            if idx >= 0:
                                if idx > search_pos:
                                    span_data.append((sp_text[search_pos:idx], False))
                                span_data.append((st, span_bold))
                                search_pos = idx + len(st)
                                matched_any = True
                    if search_pos < len(sp_text):
                        span_data.append((sp_text[search_pos:], False))
                    if not matched_any:
                        # 页面行与结构文本不匹配:整段统一加粗(按行级判断)
                        para_bold_any = any(
                            line.bold for line in page_lines
                            if line.text.strip() and len(line.text.strip()) >= 2
                        )
                        span_data = [(sp_text, para_bold_any)]
                    # 标题样式整段加粗;正文按行级加粗(span 级)
                    if style in ("title", "heading"):
                        para_bold = True
                        span_data = [(sp_text, True)]
                        out_text = sp_text
                    else:
                        para_bold = any(b for _, b in span_data)
                        out_text = "".join(t for t, _ in span_data)
                    # 校验拼接是否还原原文(近似匹配可能丢字,回退整段)
                    if out_text != sp_text:
                        span_data = [(sp_text, para_bold)]
                        out_text = sp_text
                    para.lines.append(TextLine(
                        text=out_text, bbox=(0, 0, 0, 0), size=base_size,
                        font=font_name, bold=para_bold, align="left",
                        spans=span_data,
                    ))
                    # 正文段落:匹配页面行推断首行缩进(结构树无缩进信息)
                    if style == "body":
                        sp_norm = _norm_for_match(sp_text)
                        best_x0 = 0.0
                        for line in page_lines:
                            lt_norm = _norm_for_match(line.text)
                            if not lt_norm:
                                continue
                            if lt_norm in sp_norm or sp_norm in lt_norm:
                                best_x0 = max(best_x0, line.bbox[0])
                        if best_x0 > 100:
                            para.indent_first = best_x0 - 90.0
                    add_paragraph(doc, para, font_name, base_size,
                                  pdf=pdf, image_dir=image_dir, page_width=page.rect.width,
                                  struct_path=True)
                    para_count += 1
                    if not struct_first_done:
                        struct_first_done = True
                    preview_blocks.append({
                        "type": "text",
                        "text": out_text,
                        "style": style,
                        "align": "left",
                        "spans": span_data,
                    })
        else:
            pending_para: Optional[Paragraph] = None  # 跨页待续接的未完成段落(暂未写入 docx)
            for i, page in enumerate(pdf):
                if progress_cb:
                    progress_cb(i + 1, total)
                # 可选:按 PDF 原页插入分页符(默认关闭,避免 Word 半空页)
                if i > 0 and page_breaks:
                    doc.add_page_break()
                lines, images = _extract_elements(page)
                paragraphs = rebuild_paragraphs(
                    lines, page.rect.width, page.rect.height, images,
                    first_text=not first_text_done,
                )

                # 跨页段落合并:把上一页未完成段接到本页开头文本段之前
                if pending_para is not None:
                    first_text_para = next(
                        (p for p in paragraphs if p.image is None and p.text.strip()),
                        None,
                    )
                    if first_text_para is not None:
                        # 上一页未完成段的行 + 本页首段行 = 完整段落
                        merged_lines = pending_para.lines + first_text_para.lines
                        pending_para.lines = merged_lines
                        pending_para.bold = _infer_para_bold(merged_lines)
                        paragraphs.remove(first_text_para)
                        paragraphs.insert(0, pending_para)
                    else:
                        # 本页无文本(纯图片页):pending 无法续接,直接写入
                        paragraphs.insert(0, pending_para)
                    pending_para = None

                # 确定本页最后一个未完成段落(若存在则不写入,留作跨页续接)
                last_incomplete = None
                for para in reversed(paragraphs):
                    if para.image is None and para.text.strip() and not _is_para_complete(para):
                        last_incomplete = para
                        break

                # 写入本页所有段落(除 last_incomplete 外)
                for para in paragraphs:
                    if para is last_incomplete:
                        continue  # 暂不写入,待下页续接
                    img_path = add_paragraph(
                        doc, para, font_name, base_size,
                        pdf=pdf, image_dir=image_dir, page_width=page.rect.width,
                    )
                    if para.image is not None:
                        if img_path:
                            image_count += 1
                        preview_blocks.append({
                            "type": "image",
                            "path": img_path,
                        })
                    else:
                        para_count += 1
                        if not first_text_done:
                            first_text_done = True
                        preview_blocks.append({
                            "type": "text",
                            "text": para.text,
                            "style": para.style,
                            "align": para.align,
                            # span 级加粗:每行各 span 的 (文本, 加粗) 列表(预览用)
                            "spans": [(t, b) for line in para.lines for t, b in line.spans]
                            if para.style == "body" else [],
                        })

                # 下页待续接
                pending_para = last_incomplete

            # 处理末尾残留的未完成段落(文档结尾无下一页)
            if pending_para is not None:
                img_path = add_paragraph(
                    doc, pending_para, font_name, base_size, pdf=None,
                    image_dir=image_dir, page_width=595.0,
                )
                if pending_para.image is not None:
                    if img_path:
                        image_count += 1
                    preview_blocks.append({"type": "image", "path": img_path})
                else:
                    para_count += 1
                    if not first_text_done:
                        first_text_done = True
                    preview_blocks.append({
                        "type": "text",
                        "text": pending_para.text,
                        "style": pending_para.style,
                        "align": pending_para.align,
                        "spans": [(t, b) for line in pending_para.lines for t, b in line.spans]
                        if pending_para.style == "body" else [],
                    })
    finally:
        pdf.close()

    doc.save(docx_path)

    return {
        "pages": total,
        "paragraphs": para_count,
        "images": image_count,
        "output": docx_path,
        "font": font_name,
        "font_size": font_size_label,
        "preview": preview_blocks,
    }
