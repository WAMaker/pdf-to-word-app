# -*- coding: utf-8 -*-
"""测试:生成中文测试 PDF 并转换,验证段落重建效果"""
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import fitz
from docx import Document

from app.converter import convert_pdf_to_docx, rebuild_paragraphs, _extract_lines

TEST_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_PDF = os.path.join(TEST_DIR, "sample.pdf")
SAMPLE_DOCX = os.path.join(TEST_DIR, "sample.docx")


def make_sample_pdf(path: str):
    """手工生成一个含标题、多段正文、首行缩进、居中文本的 PDF"""
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4

    # 标题(居中、大号、加粗)
    page.insert_text(
        (297 - 80, 80), "老人健康饮食指南", fontsize=24, fontname="china-s",
        color=(0, 0, 0), render_mode=0,
    )

    # 副标题
    page.insert_text((297 - 60, 120), "—— 给家人的一份贴心建议", fontsize=14, fontname="china-s")

    # 正文段落 1(首行缩进模拟)
    body1 = "随着年龄增长,身体机能逐渐下降,合理的饮食搭配显得尤为重要。老年人应当注重蛋白质的摄入,同时控制油脂和盐分的用量,保持营养均衡。"
    y = 170
    for i in range(0, len(body1), 24):
        seg = body1[i:i + 24]
        page.insert_text((60 + (12 if i == 0 else 0), y), seg, fontsize=12, fontname="china-s")
        y += 18

    # 正文段落 2(正常左对齐,无缩进)
    body2 = "每天坚持适量运动,如散步、打太极等,有助于增强体质。多喝水,多吃新鲜蔬菜水果,少吃辛辣刺激的食物,保证充足的睡眠。"
    y += 10
    for i in range(0, len(body2), 24):
        seg = body2[i:i + 24]
        page.insert_text((60, y), seg, fontsize=12, fontname="china-s")
        y += 18

    # 居中文本(模拟引用/说明)
    page.insert_text((297 - 90, y + 30), "—— 以上建议仅供参考 ——", fontsize=12, fontname="china-s")

    doc.save(path)
    doc.close()
    print(f"[OK] 测试 PDF 已生成: {path}")


def verify_docx(path: str):
    """检查 docx 段落结构:验证是否重建为逻辑段落"""
    doc = Document(path)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    print(f"\n[INFO] docx 共 {len(paras)} 个非空段落:")
    for i, p in enumerate(paras):
        style = p.style.name if p.style else "?"
        align = p.alignment
        indent = p.paragraph_format.first_line_indent
        print(f"  段{i}: style={style} align={align} indent={indent} text={p.text[:30]}...")
    return paras


def main():
    make_sample_pdf(SAMPLE_PDF)

    # 先看提取的原始行数
    pdf = fitz.open(SAMPLE_PDF)
    lines = _extract_lines(pdf[0])
    pdf.close()
    print(f"[INFO] 提取原始行数: {len(lines)}")
    for l in lines:
        print(f"  行: size={l.size:.0f} bbox={[round(v) for v in l.bbox]} align={l.align} text={l.text[:25]}")

    # 段落重建
    paras = rebuild_paragraphs(lines, 595, 842)
    print(f"\n[INFO] 重建段落数: {len(paras)} (期望约 4 段)")
    for i, p in enumerate(paras):
        print(f"  段{i}: style={p.style} align={p.align} indent_first={p.indent_first:.0f} lines={len(p.lines)} text={p.text[:30]}")

    # 完整转换
    result = convert_pdf_to_docx(SAMPLE_PDF, SAMPLE_DOCX, font_size_label="三号")
    print(f"\n[INFO] 转换结果: {result}")
    verify_docx(SAMPLE_DOCX)

    # 基本断言
    assert len(paras) >= 3, f"段落重建异常: 只有 {len(paras)} 段"
    assert any(p.style == "title" for p in paras), "未识别标题"
    print("\n✅ 测试通过!")


def test_images():
    """验证图片提取与嵌入"""
    import tempfile
    # 生成含图片的测试 PDF(插入真正的位图)
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    # 生成一张红色位图
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 200, 100))
    pix.set_rect(fitz.IRect(0, 0, 200, 100), (204, 30, 30))  # 红色填充(0-255)
    page.insert_image(fitz.Rect(60, 100, 300, 200), pixmap=pix)
    page.insert_text((60, 260), "图片下方的文字段落", fontsize=12, fontname="china-s")
    pdf_path = os.path.join(TEST_DIR, "sample_img.pdf")
    doc.save(pdf_path)
    doc.close()

    img_dir = tempfile.mkdtemp()
    result = convert_pdf_to_docx(pdf_path, SAMPLE_DOCX, font_size_label="三号", image_dir=img_dir)
    assert result["images"] >= 1, "未提取到图片"
    blocks = result["preview"]
    img_blocks = [b for b in blocks if b["type"] == "image"]
    assert img_blocks, "预览数据中没有图片块"
    assert any(os.path.isfile(b.get("path", "")) for b in img_blocks), "图片缓存文件不存在"

    # 验证 docx 内嵌图片
    docx = Document(SAMPLE_DOCX)
    assert len(docx.inline_shapes) >= 1, "docx 未嵌入图片"
    print(f"\n✅ 图片测试通过: 提取 {result['images']} 张,docx 内嵌 {len(docx.inline_shapes)} 张")


def test_realistic_docx():
    """回归测试:使用真实排版风格的经络课件 PDF(多页、跨页段落、列表标题、插图)
    覆盖修复的历史 bug:
    - 全文被误判为居中对齐(左右边距均>10pt 的行不应都判 center)
    - 段落碎片化(每行一个 block 时 block_id 判断导致同段被拆开)
    - 多页文档无分页符
    - 每页第一个加粗段落被误判为 title(应只有全文档第一个)
    - 列表符标题('l xxx')与正文合并
    """
    import tempfile
    pdf_path = os.path.join(os.path.dirname(TEST_DIR), "samples", "抱朴-經絡輔導課20.pdf")
    if not os.path.isfile(pdf_path):
        print("\n⚠️ 经络课件样例不存在,跳过真实排版回归测试")
        return
    docx_path = os.path.join(TEST_DIR, "sample_realistic.docx")
    img_dir = tempfile.mkdtemp()
    result = convert_pdf_to_docx(pdf_path, docx_path, font_size_label="三号", image_dir=img_dir)

    # 1) 结构树模式:样式忠实 PDF 结构化信息
    # 段落边界来自结构树,加粗来自字体名(Bold 字体);不放大标题字号
    # 课程标题'經絡輔導課...'应为独立段落且加粗
    d0 = Document(docx_path)
    title_paras = [p for p in d0.paragraphs if p.text.strip().startswith("經絡輔導課")]
    assert title_paras, "课程标题段落缺失"
    tp = title_paras[0]
    assert any(r.font.bold for r in tp.runs if r.text.strip()), "课程标题应加粗(字体名驱动)"
    sizes = {r.font.size.pt for r in tp.runs if r.text.strip() and r.font.size}
    assert sizes == {16.0}, f"课程标题字号应=正文16pt(结构树模式不放大),实际{sizes}"

    # 2) 正文对齐:非居中文本不应被误判为 center
    body = [b for b in result["preview"] if b["type"] == "text" and b["style"] == "body"]
    center_bodies = [b for b in body if b["align"] == "center"]
    assert not center_bodies, f"正文被误判为居中: {[b['text'][:20] for b in center_bodies]}"

    # 3) 段落重建:应合并为完整逻辑段落(不是每行一段)
    assert result["paragraphs"] < 150, f"段落仍碎片化: {result['paragraphs']} 段(21页应~100段)"
    long_bodies = [b for b in body if len(b["text"]) > 80]
    assert long_bodies, "正文段落未合并,仍然碎片化"

    # 4) 页面应为 A4,且默认流式分页(无强制分页符,避免 Word 半空页)
    import zipfile
    d = Document(docx_path)
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert xml.count('w:br w:type="page"') == 0, "默认不应有强制分页符(流式分页)"
    sec = d.sections[0]
    assert abs(sec.page_width.mm - 210) < 1 and abs(sec.page_height.mm - 297) < 1, \
        f"页面应为 A4,实际 {sec.page_width.mm:.1f}x{sec.page_height.mm:.1f}mm"

    # 4b) 开启 page_breaks 时应有分页符(21 页 → 20 个)
    docx_pb = docx_path.replace(".docx", "_pb.docx")
    convert_pdf_to_docx(pdf_path, docx_pb, font_size_label="三号", image_dir=img_dir, page_breaks=True)
    with zipfile.ZipFile(docx_pb) as z:
        xml_pb = z.read("word/document.xml").decode("utf-8")
    assert xml_pb.count('w:br w:type="page"') == 20, "page_breaks=True 时应有 20 个分页符"
    os.remove(docx_pb)

    # 5) 插图保留
    assert result["images"] == 1, f"应保留 1 张插图,实际 {result['images']}"

    # 6) 缩进:左缩进应为相对页面边距的额外缩进(PDF 左边距不应写入 left_indent)
    body_paras = [p for p in d.paragraphs if p.text.strip() and p.paragraph_format.first_line_indent]
    assert body_paras, "未检测到首行缩进段落"
    for p in body_paras[:5]:
        li = p.paragraph_format.left_indent
        assert li is None or li.pt < 10, f"左缩进异常(应≈0): {li.pt if li else None}pt"

    # 7) 加粗保留:PDF 中加粗段落应写入 docx run 加粗(结构树模式整段布尔)
    bold_paras = [p for p in d.paragraphs if p.text.strip()
                  and any(r.font.bold for r in p.runs if r.text.strip())]
    assert len(bold_paras) >= 5, f"加粗段落过少: {len(bold_paras)}(期望≥5,PDF原文大量加粗)"

    # 8) 结构树模式:段落边界来自官方结构树,不要求行级混合加粗
    # (启发式模式才做行级/span级加粗;结构树模式段落即作者原段)

    # 9) 段落边界:结构树给出作者真实分段
    # 注意:结构树模式不合并列举引导语(作者在 Word 里就是两段)
    all_text = "\n".join(p.text for p in d.paragraphs if p.text.strip())

    # 10) 跨页段落合并:PDF 分页截断的句子应续接为完整段
    assert "再來看胃經，就能夠得出" in all_text.replace("\n", ""), "跨页段落未合并(经络PDF第1-2页)"

    # 11) 标题完整性:无标点的短标题不应被跨页逻辑吞并
    assert "胃經腹部重點穴位診斷作用及功能複習\n" in all_text + "\n", \
        "短标题被跨页合并逻辑吞并"
    print(f"\n✅ 真实排版回归测试通过: {result['paragraphs']} 段, {result['images']} 图, 加粗 {len(bold_paras)} 段")


def test_ear_docx():
    """第二份样例:耳诊课件 PDF(18页、6图、列表引导语'定一區：'等)
    覆盖:跨页合并、列表引导语跨行合并、图片提取。
    """
    import tempfile
    pdf_path = os.path.join(os.path.dirname(TEST_DIR), "samples", "抱樸-耳診輔導課2.pdf")
    if not os.path.isfile(pdf_path):
        print("\n⚠️ 耳诊样例不存在,跳过")
        return
    docx_path = os.path.join(TEST_DIR, "sample_ear.docx")
    img_dir = tempfile.mkdtemp()
    result = convert_pdf_to_docx(pdf_path, docx_path, font_size_label="三号", image_dir=img_dir)
    assert result["pages"] == 18, f"页数不对: {result['pages']}"
    assert result["images"] == 6, f"图片数不对: {result['images']}"
    d = Document(docx_path)
    all_text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
    # 结构树模式:段落边界 = 作者真实分段
    # '定二到四區：...再找到四區邊界' 在作者原文中是同一大段(不强制拆分)
    assert "再找到四區邊界" in all_text.replace("\n", ""), "列表引导语内容缺失"
    # 完整独立列表项保持独立(作者分段的)
    assert "定六區：對耳輪上腳所連接的區域是耳輪六區。" in all_text, "完整列表项被误合并"
    # 章节标题(冒号结尾的加粗短行)应独立,不被引导语合并逻辑吞并
    assert "第一部分：耳診課內容回顧" in all_text, "章节标题'第一部分'被误合并"
    assert "第二部分作業講解" in all_text.replace(" ", ""), "章节标题'作業講解'被误合并"
    assert "第三部分問題答疑" in all_text.replace(" ", ""), "章节标题'問題答疑'被误合并"
    # 列表项(短'第X，'行)应独立成段
    assert "第一，耳舟，對應上肢；" in all_text, "列表项'第一，耳舟'未独立"
    # 选项行独立(结构树模式输出清理了空格/括号: 'A.即耳輪3區...')
    assert "即耳輪3區，三角窩前方的耳輪處" in all_text, "选项A未独立"
    assert "即耳輪4區，三角窩前方的耳輪處" in all_text, "选项B未独立"
    # 答案行独立(结构树输出'[答案]D'或'答案：D')
    assert ("答案D" in all_text or "[答案]D" in all_text) and "ABCD" in all_text, \
        "答案行未独立"
    # 选项行样式:选项/答案应为正文(16pt),不是标题
    option_paras = [x for x in d.paragraphs if x.text.strip()
                    and re.match(r"^[A-D][.、．]", x.text.strip())]
    assert option_paras, "未找到选项行"
    for x in option_paras:
        size = max((rr.font.size.pt if rr.font.size else 0) for rr in x.runs)
        assert size <= 16.5, f"选项行被误判为标题({size}pt): {x.text[:25]}"
    answer_paras = [x for x in d.paragraphs if x.text.strip()
                    and re.match(r"^[\[【]?答案", x.text.strip())]
    for x in answer_paras:
        size = max((rr.font.size.pt if rr.font.size else 0) for rr in x.runs)
        assert size <= 16.5, f"答案行被误判为标题({size}pt): {x.text[:25]}"
    os.remove(docx_path)
    print(f"\n✅ 耳诊样例测试通过: {result['paragraphs']} 段, {result['images']} 图")


def test_font_signal():
    """加粗检测多信号源:不依赖字体名(Bold/Black),FontWeight/黑体名也识别"""
    from app.converter import _font_name_is_bold, _analyze_fonts

    # 字体名信号
    assert _font_name_is_bold("SimHei"), "黑体(SimHei)应识别为加粗"
    assert _font_name_is_bold("Heiti SC"), "黑体(Heiti)应识别为加粗"
    assert not _font_name_is_bold("SimSun"), "宋体不应误判为加粗"
    assert not _font_name_is_bold("KaiTi"), "楷体不应误判为加粗"
    assert not _font_name_is_bold("MicrosoftYaHei"), "雅黑常规不应误判(含hei词根但排除yahei)"
    assert _font_name_is_bold("MicrosoftYaHei-Bold"), "雅黑粗体应识别"
    assert _font_name_is_bold("PingFangSC-Semibold"), "半粗体应识别"

    # FontWeight 信号:重命名粗体字体(去掉 Bold 后缀)仍应识别
    import shutil
    src = os.path.join(os.path.dirname(TEST_DIR), "samples", "抱朴-經絡輔導課20.pdf")
    if os.path.isfile(src):
        dst = os.path.join(TEST_DIR, "font_renamed.pdf")
        shutil.copy(src, dst)
        pdf = fitz.open(dst)
        for xref in range(1, pdf.xref_length()):
            try:
                obj = pdf.xref_object(xref, compressed=False)
            except Exception:
                continue
            if "MicrosoftYaHei-Bold" in obj:
                pdf.update_object(xref, obj.replace("MicrosoftYaHei-Bold", "MyFont-Bd"))
        pdf.save(dst.replace(".pdf", "_2.pdf"), incremental=False,
                 encryption=fitz.PDF_ENCRYPT_KEEP)
        pdf.close()
        pdf2 = fitz.open(dst.replace(".pdf", "_2.pdf"))
        fm = _analyze_fonts(pdf2)
        pdf2.close()
        assert fm.get("myfont-bd") is True, f"FontWeight信号应识别粗体: {fm}"
        assert fm.get("microsoftyahei") is False, f"常规字体不应误判: {fm}"
        os.remove(dst)
        os.remove(dst.replace(".pdf", "_2.pdf"))
    print("\n✅ 字体信号测试通过: 字体名/FontWeight/黑体名多信号源")


if __name__ == "__main__":
    main()
    test_images()
    test_realistic_docx()
    test_ear_docx()
    test_font_signal()
